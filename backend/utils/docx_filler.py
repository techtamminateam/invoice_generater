from docx import Document
from copy import deepcopy

def fill_document(template_path, output_path, data, employees=None):
    doc = Document(template_path)

    # Fill paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)
    
    # Fill table and add employee rows
    for table in doc.tables:
        # First, add employee rows if provided
        if employees:
            add_employee_rows(table, employees)
        
        # Then fill all placeholders in the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
    
    doc.save(output_path)

def add_employee_rows(table, employees):
    """
    Add multiple employee rows to the invoice table
    """
    # Find the template row with [name] placeholder
    template_row_index = None
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if "[name]" in cell.text:
                template_row_index = i
                break
        if template_row_index is not None:
            break
    
    if template_row_index is None:
        print("Template row with [name] not found")
        return
    
    template_row = table.rows[template_row_index]
    
    # Fill first employee in the template row
    fill_employee_row(template_row, 1, employees[0])
    
    # Add additional rows for remaining employees
    for idx in range(1, len(employees)):
        # Add a new row by copying the table row element
        new_row = add_row_after(table, template_row_index + idx - 1)
        fill_employee_row(new_row, idx + 1, employees[idx])

def add_row_after(table, row_index):
    """Add a new row after the specified index by cloning"""
    # Get the row to clone
    row_to_clone = table.rows[row_index]
    
    # Clone the row element
    tbl = table._tbl
    new_tr = deepcopy(row_to_clone._tr)
    
    # Insert after the specified row
    tbl.insert(row_index + 2, new_tr)  # +2 because index is 0-based but insert position is 1-based relative to row
    
    return table.rows[row_index + 1]

def fill_employee_row(row, serial_no, employee_data):
    """Fill a row with employee data by replacing text in runs"""
    cells = row.cells
    
    # Data mapping
    data_map = [
        str(serial_no),
        employee_data.get('name', ''),
        employee_data.get('doj', ''),
        employee_data.get('total_days', ''),
        employee_data.get('working_days', ''),
        employee_data.get('status', 'Active'),
        employee_data.get('location', ''),
        employee_data.get('net_amount', '')
    ]
    
    if len(cells) >= 8:
        for i, value in enumerate(data_map):
            cell = cells[i]
            # Replace text in all runs of the first paragraph
            if cell.paragraphs:
                para = cell.paragraphs[0]
                if para.runs:
                    # Clear all runs except the first
                    for run in para.runs[1:]:
                        run.text = ''
                    # Set the value in the first run
                    para.runs[0].text = value
                else:
                    # No runs exist, add one
                    para.add_run(value)

if __name__ == "__main__":
    template_path = "templates\\INR INVOICE.docx"
    output_path = "filled_invoice.docx"
    
    # Static data for placeholders
    data = {
        "[Invoice number]": "INV-001",
        "[Date]": "2024-11-03",
        "[GST number]": "27AAEPM1234C1Z5",
        "[PO number]": "PO-5678",
        "[sub_total]": "₹1,50,000",
        "[CGST]": "₹13,500",
        "[SGST]": "₹13,500",
        "[TIA]": "₹1,77,000"
    }
    
    # Dynamic employee data
    employees = [
        {
            'name': 'Chandu',
            'doj': '2024-05-15',
            'total_days': '22',
            'working_days': '22',
            'status': 'Active',
            'location': 'Hyderabad',
            'net_amount': '₹55,000'
        },
        {
            'name': 'Ravi Kumar',
            'doj': '2024-06-01',
            'total_days': '22',
            'working_days': '20',
            'status': 'Active',
            'location': 'Bangalore',
            'net_amount': '₹45,000'
        },
        {
            'name': 'Priya Sharma',
            'doj': '2024-04-10',
            'total_days': '22',
            'working_days': '22',
            'status': 'Active',
            'location': 'Mumbai',
            'net_amount': '₹50,000'
        }
    ]
    
    fill_document(template_path, output_path, data, employees)